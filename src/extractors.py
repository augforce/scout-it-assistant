"""Extract plain text from a file's raw bytes, dispatched by mime type."""

from __future__ import annotations

import io

from docx import Document as DocxDocument
from pypdf import PdfReader

from src.drive import DOCX, GOOGLE_DOC, GOOGLE_SHEET, PDF, XLSX


def extract_text(data: bytes, mime_type: str) -> str:
    if mime_type == GOOGLE_DOC:
        return data.decode("utf-8", errors="replace")
    if mime_type == GOOGLE_SHEET:
        # Drive exports Google Sheets as XLSX bytes (see drive.download_file).
        return _xlsx_to_text(data)
    if mime_type == PDF:
        return _pdf_to_text(data)
    if mime_type == DOCX:
        return _docx_to_text(data)
    if mime_type == XLSX:
        return _xlsx_to_text(data)
    return ""


def _pdf_to_text(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    parts = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            continue
    return "\n\n".join(parts)


def _docx_to_text(data: bytes) -> str:
    doc = DocxDocument(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _xlsx_to_text(data: bytes) -> str:
    # Header-aware row formatting: every non-header row is emitted as
    # "Col1: val | Col2: val | ..." so each row stays self-describing after
    # chunking splits the sheet across chunk boundaries. Sheets without a
    # detectable header (e.g. freeform two-table layouts) fall back to plain
    # tab-separated rows so the data is still indexed without bogus column tags.
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(data), data_only=True, read_only=True)
    parts: list[str] = []
    for sheet in wb.worksheets:
        rows = [
            [str(c) if c is not None else "" for c in row]
            for row in sheet.iter_rows(values_only=True)
        ]
        rows = [r for r in rows if any(c.strip() for c in r)]
        if not rows:
            continue

        parts.append(f"# Sheet: {sheet.title}")

        candidate_idx = max(
            range(min(5, len(rows))),
            key=lambda i: sum(1 for c in rows[i] if c.strip()),
        )
        candidate = rows[candidate_idx]
        non_empty = [c for c in candidate if c.strip()]
        # Header heuristic: short, digit-free, label-like cells. Excludes data
        # rows that happen to be the widest in the top of the sheet.
        looks_like_header = bool(non_empty) and all(
            len(c) <= 50 and not any(ch.isdigit() for ch in c) for c in non_empty
        )

        if looks_like_header:
            header = candidate
            parts.append("Columns: " + " | ".join(c for c in header if c.strip()))
            for r in rows[candidate_idx + 1 :]:
                pairs = [f"{k}: {v}" for k, v in zip(header, r) if k.strip() and v.strip()]
                if pairs:
                    parts.append(" | ".join(pairs))
        else:
            for r in rows:
                parts.append("\t".join(r))
    return "\n".join(parts)
